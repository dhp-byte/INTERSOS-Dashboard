"""
utils/report_generator.py — Génération de rapports de synthèse exécutive
(Word et PDF) à partir des données réelles du classeur INTERSOS Tchad.

Deux fonctions publiques :
  - build_docx_report(sheets, cover) -> bytes
  - build_pdf_report(sheets, cover)  -> bytes

Aucune donnée n'est inventée : tout provient des DataFrames déjà chargés par
utils.data_loader. Un secteur ou une colonne absente est simplement omis.
"""

from __future__ import annotations

import io
from datetime import datetime

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

import config
from utils import data_loader as dl

BRAND_BLUE = (0x1B, 0x3A, 0x6B)


def _achievement_label(pct: float) -> str:
    th = config.ACHIEVEMENT_THRESHOLDS
    if pct >= th["on_track_min"]:
        return "En bonne voie"
    if pct >= th["watch_min"]:
        return "À surveiller"
    return "Retard critique"


def _sector_summary_rows(sheets: dict[str, pd.DataFrame]) -> list[dict]:
    rows = []
    for sheet_name, cfg in config.SECTOR_CONFIG.items():
        if sheet_name not in sheets:
            continue
        df = sheets[sheet_name]
        target_col, reached_col = cfg["target_col"], cfg["reached_col"]
        target_total = df[target_col].sum() if target_col in df.columns else None
        reached_total = df[reached_col].sum() if reached_col in df.columns else None
        pct = (100 * reached_total / target_total) if target_total else None
        rows.append(
            {
                "sector": sheet_name.replace("_", " "),
                "activities": len(df),
                "target": target_total,
                "reached": reached_total,
                "pct": pct,
            }
        )
    return rows


# ---------------------------------------------------------------------------
# Word (.docx)
# ---------------------------------------------------------------------------
def build_docx_report(sheets: dict[str, pd.DataFrame], cover: dict[str, str]) -> bytes:
    """Construit un rapport Word de synthèse exécutive et retourne ses octets."""
    doc = Document()

    # Style de base
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    ben = sheets[config.SHEET_BENEFICIARIES]
    indicators = sheets[config.SHEET_INDICATORS]

    title = doc.add_heading("INTERSOS Tchad — Rapport de synthèse MEAL / IM", level=0)
    title.runs[0].font.color.rgb = RGBColor(*BRAND_BLUE)

    p = doc.add_paragraph()
    p.add_run(f"Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}").italic = True
    if cover.get("Mission"):
        doc.add_paragraph(cover["Mission"])

    doc.add_heading("1. Indicateurs clés", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text = "Indicateur", "Valeur"
    kpis = [
        ("Bénéficiaires enregistrés", f"{len(ben):,}".replace(",", " ")),
        ("Ménages", f"{ben['Household_ID'].nunique():,}".replace(",", " ")),
        ("Provinces couvertes", str(len(dl.get_provinces(sheets)))),
        ("Secteurs actifs", str(len(dl.get_sector_sheet_names(sheets)))),
        (
            "Indicateurs en bonne voie",
            f"{(indicators['Achievement_%'] >= config.ACHIEVEMENT_THRESHOLDS['on_track_min']).sum()}/{len(indicators)}",
        ),
    ]
    for label, value in kpis:
        row = table.add_row().cells
        row[0].text, row[1].text = label, value

    doc.add_heading("2. Performance par secteur", level=1)
    sect_table = doc.add_table(rows=1, cols=5)
    sect_table.style = "Light Grid Accent 1"
    for i, h in enumerate(["Secteur", "Activités", "Cible", "Atteint", "Taux"]):
        sect_table.rows[0].cells[i].text = h
    for row_data in _sector_summary_rows(sheets):
        row = sect_table.add_row().cells
        row[0].text = row_data["sector"]
        row[1].text = str(row_data["activities"])
        row[2].text = f"{row_data['target']:,.0f}".replace(",", " ") if row_data["target"] is not None else "—"
        row[3].text = f"{row_data['reached']:,.0f}".replace(",", " ") if row_data["reached"] is not None else "—"
        row[4].text = f"{row_data['pct']:.1f}%" if row_data["pct"] is not None else "—"

    doc.add_heading("3. Suivi des indicateurs", level=1)
    ind_table = doc.add_table(rows=1, cols=4)
    ind_table.style = "Light Grid Accent 1"
    for i, h in enumerate(["Indicateur", "Secteur", "Atteint (%)", "Statut"]):
        ind_table.rows[0].cells[i].text = h
    for _, r in indicators.iterrows():
        row = ind_table.add_row().cells
        row[0].text = str(r["Indicator"])
        row[1].text = str(r["Sector"])
        row[2].text = f"{r['Achievement_%']:.1f}%"
        row[3].text = _achievement_label(r["Achievement_%"])

    mismatches = dl.province_name_mismatches(sheets)
    doc.add_heading("4. Qualité des données", level=1)
    quality = dl.compute_data_quality_report(sheets)
    q_table = doc.add_table(rows=1, cols=len(quality.columns))
    q_table.style = "Light Grid Accent 1"
    for i, col in enumerate(quality.columns):
        q_table.rows[0].cells[i].text = str(col)
    for _, r in quality.iterrows():
        row = q_table.add_row().cells
        for i, val in enumerate(r):
            row[i].text = str(val)
    if mismatches["provinces_sans_geometrie"]:
        doc.add_paragraph(
            f"⚠ Provinces sans géométrie GeoJSON correspondante : {', '.join(mismatches['provinces_sans_geometrie'])}"
        )
    else:
        doc.add_paragraph("Toutes les provinces des données possèdent une géométrie GeoJSON correspondante.")

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Rapport généré automatiquement — Plateforme MEAL/IM INTERSOS Tchad")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0x8B, 0x95, 0xAB)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------
def build_pdf_report(sheets: dict[str, pd.DataFrame], cover: dict[str, str]) -> bytes:
    """Construit un rapport PDF de synthèse exécutive et retourne ses octets."""
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=1.5 * cm, bottomMargin=1.5 * cm)
    styles = getSampleStyleSheet()
    brand = ParagraphStyle(
        "Brand", parent=styles["Title"], textColor=colors.HexColor("#1B3A6B"), fontSize=18
    )
    story = [
        Paragraph("INTERSOS Tchad — Rapport de synthèse MEAL / IM", brand),
        Paragraph(f"Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}", styles["Italic"]),
        Spacer(1, 10),
    ]
    if cover.get("Mission"):
        story.append(Paragraph(cover["Mission"], styles["Normal"]))
        story.append(Spacer(1, 12))

    ben = sheets[config.SHEET_BENEFICIARIES]
    indicators = sheets[config.SHEET_INDICATORS]

    story.append(Paragraph("1. Indicateurs clés", styles["Heading2"]))
    kpi_data = [
        ["Indicateur", "Valeur"],
        ["Bénéficiaires enregistrés", f"{len(ben):,}".replace(",", " ")],
        ["Ménages", f"{ben['Household_ID'].nunique():,}".replace(",", " ")],
        ["Provinces couvertes", str(len(dl.get_provinces(sheets)))],
        ["Secteurs actifs", str(len(dl.get_sector_sheet_names(sheets)))],
    ]
    story.append(_styled_table(kpi_data, col_widths=[9 * cm, 6 * cm]))
    story.append(Spacer(1, 14))

    story.append(Paragraph("2. Performance par secteur", styles["Heading2"]))
    sect_rows = [["Secteur", "Activités", "Cible", "Atteint", "Taux"]]
    for row_data in _sector_summary_rows(sheets):
        sect_rows.append(
            [
                row_data["sector"],
                str(row_data["activities"]),
                f"{row_data['target']:,.0f}".replace(",", " ") if row_data["target"] is not None else "—",
                f"{row_data['reached']:,.0f}".replace(",", " ") if row_data["reached"] is not None else "—",
                f"{row_data['pct']:.1f}%" if row_data["pct"] is not None else "—",
            ]
        )
    story.append(_styled_table(sect_rows, col_widths=[4.5 * cm, 2.7 * cm, 2.7 * cm, 2.7 * cm, 2.4 * cm]))
    story.append(Spacer(1, 14))

    story.append(Paragraph("3. Suivi des indicateurs", styles["Heading2"]))
    ind_rows = [["Indicateur", "Secteur", "Atteint", "Statut"]]
    for _, r in indicators.iterrows():
        ind_rows.append([str(r["Indicator"]), str(r["Sector"]), f"{r['Achievement_%']:.1f}%", _achievement_label(r["Achievement_%"])])
    story.append(_styled_table(ind_rows, col_widths=[7 * cm, 3.5 * cm, 2 * cm, 3 * cm], small=True))

    doc.build(story)
    return buf.getvalue()


def _styled_table(data: list[list[str]], col_widths: list[float], small: bool = False) -> Table:
    t = Table(data, colWidths=col_widths, repeatRows=1)
    font_size = 7.5 if small else 9
    t.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1B3A6B")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTSIZE", (0, 0), (-1, -1), font_size),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#CCCCCC")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F2F4F8")]),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    return t
